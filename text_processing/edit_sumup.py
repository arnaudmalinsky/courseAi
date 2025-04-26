import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path
import re


def add_markdown_paragraph(doc, text, bullet_prefix=""):
    paragraph = doc.add_paragraph()
    if bullet_prefix:
        # Add the bullet manually as a first run (not touching the markdown part)
        paragraph.add_run(bullet_prefix)

    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])  # remove the **
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])  # remove the *
            run.italic = True
        else:
            run = paragraph.add_run(part)
    return paragraph


def edit_sumup(
        sumup_excel_path,
        law_ref_excel_path,
        course_name
):
    # Load Excel data
    # text_excel_path = "C:/Users/bossa/Documents/courseai_data/april/excel_w_llm/sumup_CMP_lecon1_1745452201.615262.xlsx"  # replace with your actual file
    texts_df = pd.read_excel(sumup_excel_path)

    # law_reference_excel_path = "C:/Users/bossa/Documents/courseai_data/april/excel_w_llm/llm_result_CMP_lecon1_1745144284.06419.xlsx"  # replace with your actual file
    law_ref_df = pd.read_excel(law_ref_excel_path)

    merged_df = law_ref_df.merge(
        texts_df[["unique_index","Filename","Index","text_sumup"]],
        on=["unique_index","Filename", "Index"], 
        how='left'
        )

    # Filter out rows where flag_law is not True
    # merged_df = merged_df[merged_df['flag_law'] == True]

    # Fill NaNs to avoid groupby issues
    merged_df.fillna('', inplace=True)

    # Create Word document
    doc = Document()
    doc.add_heading(course_name, level=0)

    # Sort to ensure correct agenda order
    merged_df.sort_values(by=['Title Context', 'Title lvl2', 'Title lvl3'], inplace=True)

    # Track previous levels to avoid repetition
    prev_lvl1 = ""
    prev_lvl2 = ""

    # Group by agenda levels
    grouped = merged_df.groupby(['Title Context', 'Title lvl2', 'Title lvl3'])
    for (lvl1, lvl2, lvl3), group in grouped:
        if lvl1 != prev_lvl1:
            doc.add_heading(str(lvl1), level=1)
            prev_lvl1 = lvl1
            prev_lvl2 = "" 
        if lvl2 != prev_lvl2:
            doc.add_heading(str(lvl2), level=2)
            prev_lvl2 = lvl2

        doc.add_heading(str(lvl3), level=3)

        first_sumup = group.iloc[0]['text_sumup']

        if len(first_sumup)!=0:
            print(first_sumup)
            
            
        # Split the input text into lines
            lines = first_sumup.splitlines()

            # Process each line
            for line in lines:
                leading_spaces = len(line) - len(line.lstrip(' '))
                
                clean_text = line.strip()
                has_bullet = clean_text.startswith("- ")
                if has_bullet:
                    clean_text = clean_text[2:]  # Remove the "-" for parsing, we add it separately

                # Determine indentation and bullet prefix
                if leading_spaces == 0:
                    # Top-level headline
                    p = add_markdown_paragraph(doc, clean_text)  # no bullet
                    p.paragraph_format.space_after = Pt(10)
                elif leading_spaces == 2:
                    p = add_markdown_paragraph(doc, clean_text, bullet_prefix="- ")
                    p.paragraph_format.left_indent = Pt(40)
                elif leading_spaces >= 4:
                    p = add_markdown_paragraph(doc, clean_text, bullet_prefix="- ")
                    p.paragraph_format.left_indent = Pt(60)
                else:
                    # Default first-level bullet
                    p = add_markdown_paragraph(doc, clean_text, bullet_prefix="- ")
                    p.paragraph_format.left_indent = Pt(20)

                    doc.add_paragraph("Articles juridiques trouvés :")
                    for _, row in group.iterrows():
                        if row["flag_law"]==True:
                            law_text = f"• {row['label']}, description : {row['law_description']}"# ({row['law_type']}, {row['institution']}, {row['corpus']}, {row['location']}, {row['date'] if pd.notna(row['date']) else 'No Date'})"
                            para = doc.add_paragraph(law_text)
                            para.style.font.size = Pt(10)

    # Save the document
    # Example
    path = Path(sumup_excel_path)

    output_path = f"{course_name}_sumup.docx"

    final_path = path.parent / output_path

    doc.save(final_path)

    print(f"Document saved to {final_path}")
